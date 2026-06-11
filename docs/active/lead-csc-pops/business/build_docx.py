"""
Gera Campanha-Leads-Preventivos-Pos-Vendas-Inova.docx com:
- Cores #FFC20E (acento) e #000000 (headings) com variações tonais
- Quebras de página nas seções certas
- Sem rodapés "Documento elaborado..."
- Sem metadados de header do browser
"""
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import re

BASE = Path(__file__).parent
OUTPUT = BASE / "Campanha-Leads-Preventivos-Pos-Vendas-Inova.docx"

# Paleta de cores
GOLD         = RGBColor(0xFF, 0xC2, 0x0E)   # #FFC20E — acento principal
GOLD_DARK    = RGBColor(0xB8, 0x86, 0x0B)   # #B8860B — h3 e subtítulos
GOLD_BG      = RGBColor(0xFF, 0xFB, 0xF0)   # fundo blockquote
BLACK        = RGBColor(0x00, 0x00, 0x00)   # headings principais
DARK_GRAY    = RGBColor(0x1A, 0x1A, 0x1A)   # corpo de texto
MED_GRAY     = RGBColor(0x55, 0x55, 0x55)   # texto secundário
TABLE_HDR_BG = RGBColor(0x00, 0x00, 0x00)   # fundo th
TABLE_HDR_FG = RGBColor(0xFF, 0xC2, 0x0E)   # texto th
TABLE_ALT    = RGBColor(0xF9, 0xF6, 0xEF)   # zebra table
BORDER_COLOR = RGBColor(0xE0, 0xD5, 0xB0)   # hr e borda table


def set_cell_bg(cell, rgb: RGBColor):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    hex_color = f"{rgb[0]:02X}{rgb[1]:02X}{rgb[2]:02X}"
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def add_page_break(doc: Document):
    p = doc.add_paragraph()
    run = p.add_run()
    run.add_break(docx_break_type())
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)


def docx_break_type():
    from docx.enum.text import WD_BREAK
    return WD_BREAK.PAGE


def set_para_border_left(para, rgb: RGBColor, size_pt=4):
    """Borda esquerda colorida (simula blockquote)."""
    pPr = para._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    hex_c = f"{rgb[0]:02X}{rgb[1]:02X}{rgb[2]:02X}"
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), str(size_pt * 8))
    left.set(qn("w:space"), "12")
    left.set(qn("w:color"), hex_c)
    pBdr.append(left)
    pPr.append(pBdr)


def set_heading_border_left(para, rgb: RGBColor):
    """Borda esquerda no heading h2."""
    set_para_border_left(para, rgb, size_pt=4)


def style_h1(doc: Document, text: str, page_break_before=False):
    if page_break_before:
        add_page_break(doc)
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(20)
    run.font.color.rgb = BLACK
    # underline dourado
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "18")
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), f"{GOLD[0]:02X}{GOLD[1]:02X}{GOLD[2]:02X}")
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p


def style_h2(doc: Document, text: str, page_break_before=False):
    if page_break_before:
        add_page_break(doc)
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.left_indent = Cm(0.4)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(13)
    run.font.color.rgb = BLACK
    set_heading_border_left(p, GOLD)
    return p


def style_h3(doc: Document, text: str):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = GOLD_DARK
    return p


def style_body(doc: Document, text: str, italic=False, indent=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(4)
    if indent:
        p.paragraph_format.left_indent = Cm(0.8)
    # Processa **bold** inline
    parts = re.split(r"(\*\*.*?\*\*)", text)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            r = p.add_run(part[2:-2])
            r.bold = True
            r.font.color.rgb = BLACK
            r.font.size = Pt(11)
        else:
            r = p.add_run(part)
            r.italic = italic
            r.font.color.rgb = DARK_GRAY
            r.font.size = Pt(11)
    return p


def style_blockquote(doc: Document, text: str):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.8)
    p.paragraph_format.right_indent = Cm(0.4)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    # Fundo dourado claro
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), "FFFBF0")
    pPr.append(shd)
    set_para_border_left(p, GOLD, size_pt=4)
    # Processa **bold** inline
    parts = re.split(r"(\*\*.*?\*\*)", text)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            r = p.add_run(part[2:-2])
            r.bold = True
            r.italic = True
            r.font.color.rgb = GOLD_DARK
            r.font.size = Pt(10.5)
        else:
            r = p.add_run(part)
            r.italic = True
            r.font.color.rgb = RGBColor(0x4A, 0x38, 0x00)
            r.font.size = Pt(10.5)
    return p


def style_bullet(doc: Document, text: str, level=0):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Cm(0.8 + level * 0.5)
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(2)
    parts = re.split(r"(\*\*.*?\*\*)", text)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            r = p.add_run(part[2:-2])
            r.bold = True
            r.font.color.rgb = BLACK
            r.font.size = Pt(11)
        else:
            r = p.add_run(part)
            r.font.color.rgb = DARK_GRAY
            r.font.size = Pt(11)
    return p


def style_numbered(doc: Document, text: str):
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.left_indent = Cm(0.8)
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(2)
    parts = re.split(r"(\*\*.*?\*\*)", text)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            r = p.add_run(part[2:-2])
            r.bold = True
            r.font.color.rgb = BLACK
            r.font.size = Pt(11)
        else:
            r = p.add_run(part)
            r.font.color.rgb = DARK_GRAY
            r.font.size = Pt(11)
    return p


def add_table(doc: Document, headers: list, rows: list):
    n_cols = len(headers)
    table = doc.add_table(rows=1 + len(rows), cols=n_cols)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT

    # Header row
    hdr_row = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr_row.cells[i]
        set_cell_bg(cell, TABLE_HDR_BG)
        p = cell.paragraphs[0]
        run = p.add_run(h)
        run.bold = True
        run.font.color.rgb = TABLE_HDR_FG
        run.font.size = Pt(10)

    # Data rows
    for ri, row_data in enumerate(rows):
        row = table.rows[ri + 1]
        bg = TABLE_ALT if ri % 2 == 1 else None
        for ci, cell_text in enumerate(row_data):
            cell = row.cells[ci]
            if bg:
                set_cell_bg(cell, bg)
            p = cell.paragraphs[0]
            parts = re.split(r"(\*\*.*?\*\*)", str(cell_text))
            for part in parts:
                if part.startswith("**") and part.endswith("**"):
                    r = p.add_run(part[2:-2])
                    r.bold = True
                    r.font.color.rgb = BLACK
                    r.font.size = Pt(10)
                else:
                    r = p.add_run(part)
                    r.font.color.rgb = DARK_GRAY
                    r.font.size = Pt(10)

    doc.add_paragraph()  # espaço após tabela


def add_separator(doc: Document):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "4")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "E0D5B0")
    pBdr.append(bottom)
    pPr.append(pBdr)


def build():
    doc = Document()

    # Margens
    for section in doc.sections:
        section.top_margin = Cm(2.2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2)

    # ── BRD ──────────────────────────────────────────────────────────────────
    style_h1(doc, "BRD — Campanha de Leads Preventivos de Peças")

    p = doc.add_paragraph()
    r = p.add_run("Inova Máquinas | FPS & Material Rodante")
    r.bold = True
    r.font.size = Pt(12)
    r.font.color.rgb = DARK_GRAY

    style_blockquote(doc,
        "**Data:** 28/05/2026  |  **Versão:** 1.0\n"
        "**Patrocinador:** Roberto Reis (Gerente de Pós-Venda)\n"
        "**Elaborado por:** Victor Bernardi"
    )

    # 1. Resumo Executivo
    style_h2(doc, "1. Resumo Executivo")
    style_body(doc,
        "A Inova Máquinas opera uma frota ativa de máquinas John Deere cujos consumidores "
        "críticos de peças — **Ferramentas de Penetração de Solo (FPS)** e **Material Rodante** "
        "— se desgastam em função das horas trabalhadas, não do calendário. Hoje, o contato "
        "comercial para reposição acontece de forma reativa (o cliente liga quando a peça quebra), "
        "o que gera perda de receita, risco de parada de equipamento e perda de timing do discurso de venda."
    )
    style_body(doc,
        "Este projeto entrega um **motor automatizado de geração de leads preventivos**: com base "
        "nos horímetros lidos diariamente de cada máquina, o sistema identifica automaticamente os "
        "ativos que estão próximos do ponto de desgaste e gera uma lista estruturada de oportunidades "
        "para os consultores e CSAs abordarem de forma proativa."
    )
    style_body(doc,
        "O resultado esperado é transformar peças de peças de um processo reativo para um "
        "**processo previsível e auditável**."
    )

    # 2. Objetivos SMART
    style_h2(doc, "2. Objetivos de Negócio (SMART)")
    add_table(doc,
        ["#", "Objetivo", "Meta", "Prazo"],
        [
            ["1", "Aumentar a taxa de contato preventivo com clientes de FPS e Rodante",
             "100% dos alerts gerados com consultor designado", "Imediato (semana 1)"],
            ["2", "Reduzir o tempo médio de tratativa do lead (Aging)",
             "Primeiro contato em até 3 dias úteis do alerta", "30 dias"],
            ["3", "Garantir auditoria sistêmica das vendas declaradas",
             "100% das \"Vendas\" cruzadas com proposta real no Protheus", "Contínuo"],
            ["4", "Gerar visibilidade financeira do pipeline de peças",
             "KPI de pipeline em R$ disponível no Daily Report", "Semana 1"],
        ]
    )

    # 3. Problema
    style_h2(doc, "3. Problema a Resolver")
    style_h3(doc, "Situação atual:")
    for bullet in [
        "Os consultores não têm visibilidade de quais máquinas da base estão próximas do limiar de desgaste de FPS ou Rodante.",
        "O contato acontece quando o cliente percebe o problema — geralmente tarde demais para uma venda de valor.",
        "Não há como auditar se o consultor realmente fez o contato e gerou uma proposta, criando risco de falso reporte de resultados.",
    ]:
        style_bullet(doc, bullet)
    style_h3(doc, "Impacto:")
    for bullet in [
        "Perda de oportunidades de venda de peças de alto giro.",
        "Risco de o cliente comprar de concorrente por falta de abordagem preventiva.",
        "Relatórios de desempenho sem validação cruzada com o ERP.",
    ]:
        style_bullet(doc, bullet)

    # 4. Solução — nova página
    style_h2(doc, "4. Solução Proposta", page_break_before=True)
    style_h3(doc, "Como funciona para o negócio")
    for i, item in enumerate([
        "**Monitoramento diário de horímetros:** Os horímetros de todos os equipamentos da base ativa são lidos e processados diariamente pelo sistema.",
        "**Publicação semanal da planilha de leads:** Toda segunda-feira, os alertas acumulados são publicados na planilha online compartilhada no OneDrive, com as colunas do lead bloqueadas para proteção dos dados.",
        "**Consultores e CSAs registram o contato:** As únicas células editáveis são o status do contato (Venda, Venda Perdida, Sem Contato) e o campo de observações livres.",
        "**Daily Report por e-mail:** A gestão recebe diariamente um painel de KPIs da campanha com a saúde comercial da equipe.",
        '**Auditoria automática via Protheus:** O sistema cruza automaticamente os leads marcados como "Venda" com as propostas reais geradas no ERP — eliminando falsos positivos.',
    ], 1):
        style_numbered(doc, item)

    style_h3(doc, "Réguas de Alerta (Quando o sistema dispara)")
    add_table(doc,
        ["Tipo de Peça", "Equipamento", "Gatilho"],
        [
            ["**FPS** (dentes, lâminas, pontas)", "Todos os equipamentos ativos", "A cada **+200 horas** de operação"],
            ["**Material Rodante**", "Tratores de Esteira (700J, 750J, 850J, 1050K)", "A cada **+1.500 horas** de operação"],
            ["**Material Rodante**", "Escavadeiras (130G, 130P, 160G, 160P, 180G, 200G, 200P, 210G, 210P, 350ZX, 350G)", "A cada **+3.000 horas** de operação"],
        ]
    )
    style_blockquote(doc,
        "**Reinício do ciclo:** Após o lead ser tratado (Venda ou Venda Perdida), o horímetro "
        "é zerado e o contador recomeça do zero para aquela máquina. Leads Sem Contato continuam acumulando horas."
    )

    # 5. Como os Alertas Funcionam — nova página
    style_h2(doc, "5. Como os Alertas Funcionam — Guia Completo", page_break_before=True)

    style_h3(doc, "O que é um alerta?")
    style_body(doc,
        "Um alerta é uma notificação automática gerada pelo sistema quando uma máquina da base "
        "ativa acumula horas suficientes de operação para indicar que suas peças de desgaste estão "
        "próximas do limite. É o sistema avisando o consultor: \"este cliente provavelmente vai "
        "precisar de peças em breve — ligue antes que ele precise\"."
    )

    style_h3(doc, "Por que gerar alertas preventivos?")
    style_body(doc,
        "Peças como dentes de caçamba, lâminas de trator e material rodante (esteiras, rodas) se "
        "desgastam proporcionalmente ao uso do equipamento. Quando a máquina para por quebra, o "
        "cliente perde produtividade e pode comprar a peça de qualquer fornecedor. Com o alerta "
        "preventivo, o consultor chega antes do problema — no momento em que o cliente ainda está "
        "operando e receptivo à negociação."
    )

    style_h3(doc, "Régua de Alertas — Carga Inicial (Primeiro Processamento)")
    style_body(doc,
        "Na primeira vez que o sistema processa um chassis, ele **não tem histórico** de quando a "
        "última peça foi trocada. Para não gerar alertas imediatos em toda a frota (o que sobrecarregaria "
        "a equipe), a regra de carga inicial é:"
    )
    style_blockquote(doc,
        "**O horímetro atual da máquina é registrado como o ponto de partida (marco zero).** "
        "O primeiro alerta só será gerado após acumular as horas definidas pela régua a partir desse momento."
    )
    style_h3(doc, "Exemplo prático:")
    for bullet in [
        "Uma escavadeira 200G entra no sistema com 4.200h no horímetro.",
        "O sistema registra 4.200h como marco zero.",
        "O próximo alerta de Rodante será gerado quando o horímetro atingir **7.200h** (4.200 + 3.000).",
        "O próximo alerta de FPS será gerado quando atingir **4.400h** (4.200 + 200).",
    ]:
        style_bullet(doc, bullet)
    style_body(doc, "Isso garante que a campanha começa de forma ordenada, sem gerar ruído desnecessário para a equipe comercial.")

    style_h3(doc, "Régua de Alertas — Cargas Subsequentes (Ciclo Contínuo)")
    style_body(doc, "Após a carga inicial, o sistema opera em ciclo contínuo. A cada atualização semanal (toda segunda-feira), o sistema:")
    for item in [
        "Lê o horímetro atual de cada máquina.",
        "Compara com o marco zero daquele chassis.",
        "Se a diferença atingir ou ultrapassar o limiar da régua, **gera um alerta**.",
    ]:
        style_numbered(doc, item)

    add_table(doc,
        ["Tipo de Peça", "Equipamento", "Limiar de Alerta"],
        [
            ["**FPS** — dentes, lâminas, pontas de caçamba e trator", "Toda a frota ativa", "+200 horas desde o marco zero"],
            ["**Material Rodante**", "Tratores de Esteira (700J, 750J, 850J, 1050K)", "+1.500 horas desde o marco zero"],
            ["**Material Rodante**", "Escavadeiras (130G, 130P, 160G, 160P, 180G, 200G, 200P, 210G, 210P, 350ZX, 350G)", "+3.000 horas desde o marco zero"],
        ]
    )
    style_blockquote(doc,
        "**Por que limiares diferentes para Rodante?** Tratores de esteira se locomovem continuamente "
        "— suas esteiras estão em atrito constante com o solo. Escavadeiras trabalham majoritariamente "
        "paradas, rodando as esteiras apenas nas transferências entre obras. Por isso, o desgaste do "
        "rodante é muito mais rápido nos tratores."
    )

    style_h3(doc, "O Ciclo de Vida de um Lead")
    for bullet in [
        "Quando o consultor marca **Venda** ou **Venda Perdida**: o sistema registra o horímetro daquele momento como o novo marco zero. O chassis sai da fila e só retorna após acumular as horas do próximo ciclo.",
        "Quando o consultor deixa em branco ou marca **Sem Contato**: o lead permanece visível na planilha e as horas continuam acumulando. O aging do lead começa a crescer e aparece no Daily Report como alerta para a gestão.",
    ]:
        style_bullet(doc, bullet)

    # 6. KPIs — nova página
    style_h2(doc, "6. Indicadores de Performance (KPIs) — Daily Report", page_break_before=True)
    style_body(doc, "Estes são os 5 indicadores acompanhados diariamente pela gestão no Daily Report:")
    add_table(doc,
        ["KPI", "O que mede", "Por que importa"],
        [
            ["**Adesão Comercial**", "% dos leads da semana com algum feedback registrado", "Mede o engajamento e velocidade da equipe em fazer os contatos"],
            ["**Taxa de Conversão Real**", "% de leads tratados que resultaram em venda", "Mede a eficiência do discurso comercial"],
            ["**Aderência de Propostas**", "% de \"Vendas\" declaradas com proposta real no Protheus", "Auditoria sistêmica — garante que o número é verdadeiro"],
            ["**Aging do Lead**", "Dias médios de um alerta ativo sem primeiro contato", "Alerta quando o timing preventivo está sendo perdido"],
            ["**Pipeline Financeiro**", "Soma em R$ das propostas abertas no Protheus vinculadas à campanha", "Mede o valor financeiro que a campanha está movimentando"],
        ]
    )

    style_h3(doc, "Destinatários do Daily Report")
    style_body(doc, "O e-mail diário com os KPIs é enviado automaticamente para:")
    add_table(doc,
        ["Nome", "E-mail"],
        [
            ["Pedro Sarnaglia", "pedro.sarnaglia@inovamaquinas.com"],
            ["Leandro Silva", "leandro.silva@inovamaquinas.com"],
            ["Marcelo Costa", "marcelo.costa@inovamaquinas.com"],
            ["Murilo Nunes", "murilo.nunes@inovamaquinas.com"],
            ["Luciana Borges", "luciana.borges@inovamaquinas.com"],
            ["Gabriela Rodarte", "gabriela.rodarte@inovamaquinas.com"],
            ["Roberto Reis", "roberto.reis@inovamaquinas.com"],
            ["Victor Bernardi (c/o — Engenharia de Dados)", "victor.bernardi@inovamaquinas.com"],
        ]
    )

    # 7. Responsabilidades — nova página
    style_h2(doc, "7. Matriz de Responsabilidades", page_break_before=True)
    add_table(doc,
        ["Papel", "Responsável", "O que faz na campanha"],
        [
            ["**Gerente de Pós-Venda**", "Roberto Reis", "Lidera reunião semanal de revisão de metas (terças-feiras); avalia KPIs consolidados"],
            ["**Gerentes / Coordenadores Regionais**", "Gabriela Rodarte","Pedro Sarnaglia, Leandro Silva, Marcelo Costa, Luciana Borges", "Acompanham o Daily Report; gerenciam os consultores de suas respectivas regiões (regiões a serem definidas)"],
            ["**CSA**", "Murilo Nunes", "Realiza o contato ativo com os clientes da sua carteira; registra o status na planilha"],
            ["**Consultores de Vendas**", "A definir — aguardando mapeamento do Murilo", "Realizam o contato ativo com os clientes; registram o status na planilha"],
            ["**Dados**", "Victor Bernardi", "Mantém o motor de cálculo, atualiza a planilha e dispara o Daily Report diariamente"],
        ]
    )
    style_blockquote(doc,
        "**Pendência:** A segmentação de clientes e consultores por CSA (carteira do Murilo) está em definição. "
        "Assim que o mapeamento for enviado, a planilha será configurada com o campo \"Responsável\" preenchido automaticamente."
    )

    # 8. Governança
    style_h2(doc, "8. Governança e Segurança dos Dados")
    for bullet in [
        "**Soberania de dados:** As colunas de origem do lead (Chassi, Cliente, CNPJ, Modelo, Motivo, Horímetro) são bloqueadas com senha. Consultores não podem editar, copiar ou deletar esses campos.",
        "**Acesso controlado:** Apenas as colunas de feedback comercial são editáveis pela equipe de vendas.",
        "**Auditoria mensal:** Cruzamento automático entre a planilha e o Protheus para validar a integridade dos resultados reportados.",
        "**Snapshot diário:** Antes de cada atualização, o sistema registra um snapshot do estado atual para rastreabilidade histórica.",
    ]:
        style_bullet(doc, bullet)

    # 9. Fora do Escopo
    style_h2(doc, "9. Fora do Escopo deste Documento")
    for bullet in [
        "Implementação técnica do motor ETL (detalhada na Especificação Técnica 2026-05-27-especificacao_leads_csc_pops.md).",
        "Regras do motor CEVAP de clientes inativos (projeto paralelo, base TIVAN 590 dias).",
        "Gestão de metas individuais dos consultores.",
    ]:
        style_bullet(doc, bullet)

    # ── RELATO DE IMPACTO — nova página ──────────────────────────────────────
    style_h1(doc, "Relato de Impacto — Campanha de Leads Preventivos de Peças", page_break_before=True)
    p = doc.add_paragraph()
    r = p.add_run("Inova Máquinas | FPS & Material Rodante")
    r.bold = True; r.font.size = Pt(12); r.font.color.rgb = DARK_GRAY
    style_blockquote(doc,
        "**Data:** 28/05/2026  |  **Versão:** 1.0 — Projeção de Implantação\n"
        "**Público:** Gerentes e Coordenadores Regionais\n"
        "**Elaborado por:** Engenharia de Dados — Victor Bernardi"
    )

    # 1. Descoberta Principal
    style_h2(doc, "1. A Descoberta Principal")
    style_blockquote(doc,
        "**A Inova Máquinas possui uma frota ativa de clientes cujas máquinas acumulam horas "
        "continuamente — e hoje nenhum consultor é alertado automaticamente quando o momento certo "
        "de venda preventiva chega.** Cada hora que passa sem contato é uma janela perdida para a concorrência."
    )

    # 2. Problema em Números
    style_h2(doc, "2. O Problema em Números")
    style_body(doc, "O desgaste de FPS e Material Rodante é **previsível e recorrente**. As peças não falham aleatoriamente — elas se desgastam em função das horas de operação, com intervalos conhecidos:")
    add_table(doc,
        ["Tipo", "Intervalo médio de troca", "Frequência por máquina/ano*"],
        [
            ["FPS (dentes, lâminas, pontas)", "A cada 200h", "~4 a 6 vezes ao ano"],
            ["Rodante — Tratores de Esteira", "A cada 1.500h", "~1 vez a cada 18 meses"],
            ["Rodante — Escavadeiras", "A cada 3.000h", "~1 vez a cada 3 anos"],
        ]
    )
    style_body(doc, "*Estimativa para máquinas operando ~800h/ano.", italic=True)
    style_body(doc, "Sem um sistema de alertas, essas oportunidades dependem inteiramente do cliente ligar — ou do consultor lembrar de ligar. **Ambos os cenários são imprecisos e sujeitos a esquecimento.**")

    # 3. O que a campanha entrega
    style_h2(doc, "3. O Que a Campanha Entrega")
    style_h3(doc, "Para os Consultores e CSAs")
    for bullet in [
        "Uma lista semanal, atualizada toda segunda-feira, com os chassis que atingiram o limiar de desgaste — sem precisar calcular nada manualmente.",
        "O histórico de cada contato registrado na própria planilha, acessível de qualquer lugar pelo OneDrive.",
    ]:
        style_bullet(doc, bullet)
    style_h3(doc, "Para os Gerentes e Coordenadores Regionais")
    style_body(doc, "Um e-mail diário com 5 indicadores que respondem as perguntas mais importantes:")
    for bullet in [
        "*Minha equipe está fazendo os contatos?* → **Adesão Comercial**",
        "*Os contatos estão virando venda?* → **Taxa de Conversão Real**",
        "*Os números que recebo são reais?* → **Aderência de Propostas (Ponte da Verdade)**",
        "*Estou perdendo timing?* → **Aging do Lead**",
        "*Quanto vale o que está em negociação?* → **Pipeline Financeiro em R$**",
    ]:
        style_bullet(doc, bullet)
    style_h3(doc, "Para a Gestão de Pós-Venda (Roberto)")
    for bullet in [
        "**Auditoria sistêmica automática:** toda \"Venda\" declarada na planilha é cruzada com as propostas reais do ERP Protheus — eliminando falsos relatórios de resultado.",
        "Base de dados para a **reunião semanal de terças-feiras** com evidências concretas de desempenho por consultor e por região.",
    ]:
        style_bullet(doc, bullet)

    # 4. Impacto
    style_h2(doc, "4. Impacto Esperado")
    style_h3(doc, "Receita")
    style_body(doc, "A campanha converte um processo reativo em previsível. Os ganhos são diretos:")
    for bullet in [
        "**Mais vendas no momento certo:** O consultor chega antes do problema — quando o cliente ainda está operando e receptivo.",
        "**Recorrência garantida:** Com o ciclo reiniciando automaticamente após cada tratativa, nenhuma máquina \"some\" do radar comercial.",
        "**Redução de perda para concorrência:** Peças compradas em emergência frequentemente vão para o fornecedor mais rápido, não para a Inova.",
    ]:
        style_bullet(doc, bullet)
    style_h3(doc, "Governança")
    for bullet in [
        "**Fim do autodeclaratório:** Qualquer \"Venda\" sem proposta no Protheus aparece imediatamente no KPI de Aderência — tornando o resultado auditável e confiável.",
        "**Visibilidade regional:** Cada gerente e coordenador recebe os mesmos KPIs diariamente, permitindo comparação entre regiões e filiais.",
    ]:
        style_bullet(doc, bullet)

    # 5. Recomendações
    style_h2(doc, "5. Recomendações Imediatas")
    add_table(doc,
        ["Prioridade", "Ação", "Responsável", "Prazo"],
        [
            ["Alta", "Equipe comercial acessar e testar a planilha no OneDrive antes do primeiro ciclo", "Gerentes","Coordenadores", "Semana 1"],
            ["Média", "Gerentes/Coordenadores regionais confirmarem suas regiões de atuação para segmentação do relatório", "Gerentes","Coordenadores", "Assim que receber o link"],
        ]
    )

    # 6. Próximos Marcos
    style_h2(doc, "6. Próximos Marcos")
    for bullet in [
        "**Semana 1** → Primeiro ciclo semanal de alertas publicado na planilha",
        "**Semana 1** → Primeiro Daily Report enviado para todos os destinatários",
        "**Semana 2** → Primeira reunião de terça com base em dados reais da campanha",
        "**Mês 1** → Primeira auditoria de Aderência de Propostas cruzando planilha x Protheus",
        "**Mês 2** → Primeiro benchmark de conversão por região disponível",
    ]:
        style_bullet(doc, bullet)

    doc.save(OUTPUT)
    print(f"DOCX gerado: {OUTPUT} ({OUTPUT.stat().st_size:,} bytes)")


if __name__ == "__main__":
    build()
